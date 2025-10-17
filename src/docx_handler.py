"""Word (DOCX) document handler for extraction and reintegration."""

from typing import Any

from docx import Document

from .base_handler import BaseDocumentHandler


class TextIterator:
    """Iterator for translated texts to match extraction order."""

    def __init__(self, translated_data):
        self.paragraphs = translated_data["paragraphs"]
        self.current_paragraph = 0
        self.current_text = 0

    def get_next(self):
        """Get next translated text."""
        if self.current_paragraph >= len(self.paragraphs):
            return None

        paragraph_texts = self.paragraphs[self.current_paragraph]["texts"]

        if self.current_text >= len(paragraph_texts):
            return None

        text = paragraph_texts[self.current_text]
        self.current_text += 1
        return text

    def next_paragraph(self):
        """Move to next paragraph."""
        self.current_paragraph += 1
        self.current_text = 0


def replace_runs_with_text_list(paragraph, text_iterator):
    """Replace runs with translated texts while preserving formatting.

    Args:
        paragraph: A paragraph object from python-docx
        text_iterator: TextIterator instance

    Returns:
        int: Number of runs replaced
    """
    count = 0

    for run in paragraph.runs:
        if run.text:  # Only replace runs that have text
            translated_text = text_iterator.get_next()
            if translated_text is not None:
                run.text = translated_text
                count += 1

    return count


class DOCXHandler(BaseDocumentHandler):
    """Handler for Word (DOCX) files."""

    def extract_text(self, docx_path):
        """Extract all text from a DOCX file in deterministic order.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            dict: JSON structure with paragraphs and runs
        """
        document = Document(docx_path)

        result: dict[str, list[dict[str, Any]]] = {"paragraphs": []}

        for paragraph in document.paragraphs:
            paragraph_texts = []

            # Extract each run separately to preserve formatting
            for run in paragraph.runs:
                if run.text:  # Include runs with any text
                    paragraph_texts.append(run.text)

            # Add paragraph entry (even if empty to preserve structure)
            result["paragraphs"].append({"texts": paragraph_texts})

        # Handle tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_texts = []
                        for run in paragraph.runs:
                            if run.text:
                                paragraph_texts.append(run.text)
                        result["paragraphs"].append({"texts": paragraph_texts})

        return result

    def print_extraction_summary(self, data):
        """Print extraction summary.

        Args:
            data: Extracted data dictionary
        """
        total_texts = sum(len(para["texts"]) for para in data["paragraphs"])
        print(f"✓ Extracted {total_texts} text elements from {len(data['paragraphs'])} paragraphs")

    def reintegrate_text(self, original_docx_path, translated_data, output_docx_path):
        """Reintegrate translated text into a DOCX file.

        Args:
            original_docx_path: Path to original DOCX file
            translated_data: Dictionary with translated text
            output_docx_path: Path to save output DOCX file

        Returns:
            int: Total number of text elements replaced
        """
        # Load original document
        document = Document(original_docx_path)

        text_iterator = TextIterator(translated_data)

        total_replaced = 0

        # Replace text in main document paragraphs
        for paragraph in document.paragraphs:
            count = replace_runs_with_text_list(paragraph, text_iterator)
            total_replaced += count
            text_iterator.next_paragraph()

        # Replace text in table cells
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count = replace_runs_with_text_list(paragraph, text_iterator)
                        total_replaced += count
                        text_iterator.next_paragraph()

        # Save modified document
        document.save(output_docx_path)

        return total_replaced
